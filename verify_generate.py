import os
import sqlite3
from docx import Document
from main import app, generate_paper, GenerateRequest, db, MEDIA_DIR
import asyncio

# Setup
print("--- Testing /generate Endpoint ---")

async def test_generate():
    # 1. Mock Request
    # Request 5 random questions
    req = GenerateRequest(types=[], total_count=5)
    
    try:
        # Call endpoint (it's async)
        # We can't easily await FileResponse to get the file content directly in this script 
        # without running a full server or mocking more deepy.
        # But `generate_paper` returns a FileResponse. 
        # The file is created at `output_path`.
        
        # ACTUALLY, `generate_paper` in main.py is defined as `async def`.
        # We can call it, but we need to intercept the file path before it returns FileResponse?
        # Or just check the "temp" directory for the newest file?
        
        # Better: let's invoke the logic manually to mimic the endpoint, 
        # OR just run the server and use requests? 
        # Running server is complex in this environment.
        
        # Let's inspect `main.py` again. `generate_paper` returns `FileResponse(output_path, ...)`
        # `FileResponse.path` attribute should hold the path.
        
        res = await generate_paper(req)
        print(f"Endpoint returned: {type(res)}")
        
        output_path = res.path
        print(f"Output File: {output_path}")
        
        if not os.path.exists(output_path):
            print("❌ File was not created.")
            return

        # 2. Check Database for UUID
        # We need to find the UUID. It's not returned by the API directly, it is in the file.
        # Let's read the file first.
        doc = Document(output_path)
        if len(doc.paragraphs) > 0:
            first_line = doc.paragraphs[0].text
            print(f"First Line: '{first_line}'")
            if first_line.startswith("Paper ID: "):
                paper_uuid = first_line.replace("Paper ID: ", "").strip()
                print(f"✅ Found Paper ID: {paper_uuid}")
                
                # Check DB
                conn = db.get_connection()
                c = conn.cursor()
                c.execute("SELECT uuid, question_ids FROM generated_papers WHERE uuid = ?", (paper_uuid,))
                row = c.fetchone()
                conn.close()
                
                if row:
                    print(f"✅ Found DB Record for {paper_uuid}")
                    print(f"Stored Question IDs: {row[1]}")
                else:
                    print(f"❌ DB Record mismatch for {paper_uuid}")
            else:
                print("❌ 'Paper ID' not found in first line.")
        else:
            print("❌ Document is empty.")

        # Cleanup
        # os.remove(output_path) 

    except Exception as e:
        print(f"❌ Exception: {e}")
        import traceback
        traceback.print_exc()

# Run
if __name__ == "__main__":
    asyncio.run(test_generate())
