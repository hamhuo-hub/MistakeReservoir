import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup

class PaperBuilder:
    def __init__(self, media_dir: str):
        self.media_dir = media_dir
        
    def create_paper(self, questions: list, output_path_base: str, paper_uuid: str = None):
        """
        Generates two files:
        1. Question Paper
        2. Answer Key
        Returns list of generated file paths.
        """
        doc_q = Document()
        doc_a = Document()
        
        # --- Styles Setup for Both ---
        for doc in [doc_q, doc_a]:
            style = doc.styles['Normal']
            style.font.name = 'Microsoft YaHei'
            style.font.size = Pt(10.5) # 5号
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1.0
        

        
        # --- Paper ID Header (Visible for Review) ---
        if paper_uuid:
            # Only needed on Question Paper for scanning/review
            p = doc_q.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            run = p.add_run(f"Paper ID: {paper_uuid}")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(150, 150, 150) # Grey
            doc_q.add_paragraph() # Spacing
        
        # --- Logic ---
        # Sort Order
        type_map = {
            '常识': 1, 
            '言语': 2, 
            '数量': 3, 
            '判断': 4, '图形': 4.1, '定义': 4.2, '类比': 4.3, '逻辑': 4.4,
            '资料': 5
        }
        
        questions.sort(key=lambda q: (type_map.get(q.get('type', ''), 99), q.get('material_id') or 0, q.get('original_num')))
        
        current_type = None
        last_material_id = None
        global_idx = 1
        
        section_map = {
            '常识': '第一部分 常识判断',
            '言语': '第二部分 言语理解与表达',
            '数量': '第三部分 数量关系',
            '判断': '第四部分 判断推理',
            '图形': '第四部分 判断推理', # Subtypes grouped under main
            '定义': '第四部分 判断推理',
            '类比': '第四部分 判断推理',
            '逻辑': '第四部分 判断推理',
            '资料': '第五部分 资料分析'
        }
        
        # Track main sections added to avoid repeats
        added_encounters = set()
        
        for q in questions:
            q_type = q.get('type') or "其他"
            clean_type = q_type
            
            # Map subtype to main type for header
            if q_type in ['图形', '定义', '类比', '逻辑']:
                clean_type = '判断'
                
            header_title = section_map.get(clean_type, f"部分 {clean_type}")
            
            # 1. Section Header (Questions Doc)
            if header_title not in added_encounters:
                if added_encounters: doc_q.add_paragraph() # Spacing
                h = doc_q.add_heading(header_title, level=1)
                h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                added_encounters.add(header_title)
                current_type = clean_type # Logical group
                last_material_id = None
            
            # 2. Material (Questions Doc)
            mid = q.get('material_id')
            if mid and mid != last_material_id:
                mat_content = q.get('material_content')
                if mat_content:
                    p = doc_q.add_paragraph()
                    r = p.add_run("根据以下材料，回答下列问题：")
                    r.bold = True
                    self._add_html_content(doc_q, mat_content)
                    doc_q.add_paragraph() # Spacing after material
                last_material_id = mid
            elif not mid:
                last_material_id = None

            # 3. Question Logic (Questions Doc)
            # Stem
            p = doc_q.add_paragraph()
            p.add_run(f"{global_idx}. ").bold = True
            # Bold the stem text as requested
            self._add_html_content_inline(p, q.get('content_html'), doc_q, bold=True, q_type=q.get('type'))
            
            # Options
            if q.get('options_html'):
                 self._add_options(doc_q, q.get('options_html'))
            
            global_idx += 1
            doc_q.add_paragraph() # Spacing between questions

        # --- Answer Key Doc ---
        # 1. Header
        h = doc_a.add_heading('参考答案与解析', level=1)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 2. Quick Answer Table
        self._create_answer_table(doc_a, questions)
        doc_a.add_paragraph() # Spacing

        # 3. Detailed Answers
        for i, q in enumerate(questions):
            p = doc_a.add_paragraph()
            p.add_run(f"{i+1}. ").bold = True
            
            # Answer + Analysis
            # Retrieve answer string (assuming it might be HTML or distinct field?)
            # Usually 'answer_html' contains the analysis/explanation too or just the letter?
            # Based on user data, answer_html often includes "【答案】B 【解析】..."
            
            ans = q.get('answer_html')
            if ans:
                self._add_html_content_inline(p, ans, doc_a)
            else:
                p.add_run("（暂无解析）")
            doc_a.add_paragraph() # Spacing

        # Finalize
        if questions:
            self._unify_styles(doc_q)
            self._unify_styles(doc_a)

        # Save Both
        path_q = output_path_base.replace(".docx", "_题目.docx")
        path_a = output_path_base.replace(".docx", "_答案.docx")
        
        doc_q.save(path_q)
        doc_a.save(path_a)
        
        return [path_q, path_a]

    def _create_answer_table(self, doc, questions):
        """Creates a grid of answers at the top of the answer document."""
        import re
        
        # Extract simple answers (A, B, C, D) from answer_html if possible
        # Or hopefully there is a cleaner field. If not, use regex on html.
        # Common format: <p>【答案】 A</p> or just "A"
        
        extracted = {}
        for i, q in enumerate(questions):
            txt = q.get('answer_html', '')
            # Simple regex to find the letter after 答案
            # Matches: 【答案】A or 答案：A or just A (risky)
            # Let's try matching standard patterns first
            match = re.search(r'(?:答案|Answer)[^\w]*([A-H])', txt, re.IGNORECASE)
            if match:
                extracted[i+1] = match.group(1).upper()
            else:
                # Fallback: maybe the text IS just the answer if short?
                clean = re.sub(r'<[^>]+>', '', txt).strip()
                if clean in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']:
                     extracted[i+1] = clean
                else:
                     extracted[i+1] = " "
        
        if not extracted: return

        # Create Table
        chunk_size = 5
        total = len(questions)
        # rows = ceil(total/5) * 2
        
        table = doc.add_table(rows=0, cols=chunk_size)
        table.style = 'Table Grid'
        table.autofit = False 
        
        # Set column widths? (Optional, docx automatic is usually ok for 5 cols)
        
        sorted_nums = sorted(extracted.keys())
        for i in range(0, len(sorted_nums), chunk_size):
            chunk = sorted_nums[i : i + chunk_size]
            
            row_q = table.add_row().cells
            row_a = table.add_row().cells
            
            for j, q_num in enumerate(chunk):
                row_q[j].text = str(q_num)
                row_a[j].text = extracted[q_num]
                
                # Center align
                for cell in [row_q[j], row_a[j]]:
                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _unify_styles(self, doc):
        """
        Global format sweeper.
        Enforces 'Microsoft YaHei', 10.5pt, single spacing for the whole document,
        while preserving Headings (titles).
        """
        from docx.oxml.ns import qn
        
        for paragraph in doc.paragraphs:
            # Detect Heading
            is_heading = paragraph.style.name.startswith('Heading')
            
            # 1. Paragraph Format
            pf = paragraph.paragraph_format
            if not is_heading:
                # Standard body text
                pf.line_spacing = 1.0
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.first_line_indent = None # Reset indent unless specific logic needs it
                # pf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT # Optional: force left
            else:
                # Keep heading spacing/alignment usually, or enforce specific heading style if needed
                pass

            # 2. Run Format (Font & Size)
            for run in paragraph.runs:
                font = run.font
                
                # --- Set Font Family (East Asia & ASCII) ---
                font.name = 'Microsoft YaHei'
                # For python-docx to properly set CJK font:
                rPr = run._element.get_or_add_rPr()
                # Check if rFonts exists, if not create
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = db_element = rPr.makeelement(qn('w:rFonts'))
                    rPr.append(rFonts)
                
                rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
                rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
                
                # --- Set Font Size ---
                # Headings usually have their own size in styles, but we can enforce if needed.
                # Here we strictly enforce body text size.
                if not is_heading:
                    font.size = Pt(10.5) # 5 hao
                else:
                    # Optional: Could enforce specific heading sizes here
                    pass

    def _add_html_content(self, doc, html_str):
        """Block level adder"""
        if not html_str: return
        soup = BeautifulSoup(html_str, 'html.parser')
        
        # Extract images and text
        # If div.img-container -> Image
        # If p -> text
        
        for elem in soup.find_all(['p', 'div', 'table']):
            if elem.name == 'p':
                if elem.get_text().strip():
                    doc.add_paragraph(elem.get_text().strip())
            elif elem.name == 'div' and 'img-container' in elem.get('class', []):
                self._add_image(doc, elem)
            elif elem.name == 'table':
                # Simplified table
                pass

    def _add_html_content_inline(self, paragraph, html_str, doc, bold=False, q_type=None):
        """Adds text to existing paragraph, inserting images inline or as blocks based on size"""
        if not html_str: return
        soup = BeautifulSoup(html_str, 'html.parser')
        
        # Iterate over child nodes to maintain order
        # Note: This is a simplistic traversal. Nested tags might need recursion, 
        # but usually the input HTML is flat-ish (p, img, span).
        
        # Determine the run to append to
        run = paragraph.add_run()
        if bold: run.bold = True
        
        # We process 'descendants' carefully or just iterate contents?
        # contents is strictly direct children. text might be split.
        
        # Strategy: Flatten the soup to a list of (type, content)
        # Text -> append to current run
        # Img -> check size -> append to run OR break paragraph
        
        # Simple recursive walker (flattened)
        for output in self._flatten_nodes(soup):
            type_, content = output
            if type_ == 'text':
                if content: 
                     run.add_text(content)
            elif type_ == 'img':
                # Try to insert
                self._insert_image_hybrid(doc, run, content, q_type=q_type)
                # If we broke the run for a block image, we need a NEW run for subsequent text?
                # _insert_image_hybrid might handle breaks. 
                # If it adds a break, the 'run' object is still technically valid for adding text, 
                # but might be visually weird if we don't handle it.
                # However, for simplicity, we keep appending.
                pass
                
    def _flatten_nodes(self, element):
        """Yields ('text', str) or ('img', src)"""
        if element.name == 'img':
            yield ('img', element.get('src'))
            return

        if isinstance(element, str): # NavigableString
            yield ('text', str(element))
            return

        for child in element.children:
            yield from self._flatten_nodes(child)

    def _add_image(self, doc, elem):
        img = elem.find('img')
        if img:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
            self._insert_image_hybrid(doc, run, img.get('src'))

    def _add_options(self, doc, html_str):
        """Parses options and puts them on new lines, preserving images"""
        if not html_str: return
        soup = BeautifulSoup(html_str, 'html.parser')
        
        # If wrapped in <p>, iterate Ps. 
        # Else, just parse the whole blob.
        ps = soup.find_all('p')
        
        if ps:
            for p_tag in ps:
                p = doc.add_paragraph()
                self._add_html_content_inline(p, str(p_tag), doc)
        else:
            # Fallback for unwrapped text
            p = doc.add_paragraph()
            self._add_html_content_inline(p, html_str, doc)

    def _insert_image_hybrid(self, doc, run, src, q_type=None):
        """
        Inserts image. 
        - If 'q_type' contains '图形', FORCE height=4cm.
        - Otherwise:
            - If small/icon-like: Insert into 'run' with height=Pt(11) (Inline).
            - If large: Insert as new Paragraph (Block).
        """
        try:
            from PIL import Image
        except ImportError:
            Image = None

        if not src: return
        fname = src.split('/')[-1]
        fpath = os.path.join(self.media_dir, fname)
        
        if not os.path.exists(fpath):
            print(f"DEBUG: Image missing {fpath}")
            return
            
        # Determine Sizing Strategy
        is_inline = False
        width_arg = None
        height_arg = None
        
        # Check Special Type first
        if q_type and '图形' in str(q_type):
            # Force 4cm
            # Treat as inline-ish but with specific height
            is_inline = True
            height_arg = Cm(4)
        elif Image:
            try:
                with Image.open(fpath) as img:
                    w, h = img.size
                    
                    # Revised Heuristic for "Small / Inline"
                    # User specifically wants images to match 5-hao font (~10.5pt, approx 14-20px rendered).
                    # If an image is "relatively small" (e.g. < 250px height), assume it's an inline symbol/formula and shrink it.
                    # 250px is arbitrary but covers most high-dpi small icons.
                    
                    if h < 250: 
                        is_inline = True
                        height_arg = Pt(11) # Force to 5-hao size
                    else:
                        # Large content (Chart, Screenshot)
                        is_inline = False
                        if w > 400:
                            width_arg = Inches(5.5) # Max Page Width
                        else:
                             width_arg = Inches(3.5) if w > 300 else None
            except:
                is_inline = False # Fallback to block on error
                width_arg = Inches(2.0)
        
        if is_inline:
            # Add to CURRENT run
            try:
                run.add_picture(fpath, height=height_arg)
            except Exception as e:
                 print(f"Error adding inline pic: {e}")
        else:
            # Add to NEW paragraph
            # We need to break the flow? 
            # Ideally we'd close the current run, make a new p, then resume?
            # But we are inside `_add_html_content_inline` taking a `paragraph`.
            # We can't easily "split" the paragraph passed in unless we return new context.
            # Workaround: Add to the *End* of the current paragraph via run?
            # run.add_picture() adds it at the current position. 
            # If we want a "Block" feel but are stuck in a run, we can add a break before/after?
            try:
                run.add_break()
                if width_arg:
                    run.add_picture(fpath, width=width_arg)
                else:
                    run.add_picture(fpath)
                run.add_break()
            except Exception as e:
                pass
