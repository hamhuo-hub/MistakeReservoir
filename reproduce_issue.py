
import os
import shutil
from docx import Document
from extractor import QuestionExtractor

def create_test_docx(filename):
    doc = Document()
    doc.add_paragraph("1. Test Question")
    doc.add_paragraph("A. Option A")
    doc.add_paragraph("B. Option B")
    doc.add_paragraph("【解析】Analysis text.")
    doc.add_paragraph("故")
    doc.add_paragraph("故。")
    doc.add_paragraph("End of analysis.")
    doc.save(filename)

def test_extraction():
    docx_path = "test_issue.docx"
    media_dir = "test_media"
    if os.path.exists(media_dir):
        shutil.rmtree(media_dir)
    
    create_test_docx(docx_path)
    
    extractor = QuestionExtractor(media_dir=media_dir)
    questions = extractor.extract_from_file(docx_path)
    
    for q in questions:
        print(f"Question {q['original_num']}:")
        print("Answer HTML:")
        print(q['answer_html'])
        
        if "<p>故</p>" in q['answer_html'] or "<p>故。</p>" in q['answer_html']:
            print("ISSUE REPRODUCED: '故' found in answer HTML.")
        else:
            print("SUCCESS: '故' was filtered out.")

    # Cleanup
    if os.path.exists(docx_path):
        os.remove(docx_path)
    if os.path.exists(media_dir):
        shutil.rmtree(media_dir)

if __name__ == "__main__":
    test_extraction()
