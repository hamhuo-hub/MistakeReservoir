import os
import re
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

try:
    from util.converter_rules import (
        START_KEYWORD_REGEX,
        END_KEYWORD_PATTERNS,
        STRONG_DELETE_CONTAIN,
        FORCE_DELETE_PREFIXES,
        ANSWER_VAL_PATTERN
    )
except ImportError:
    import sys
    sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    from util.converter_rules import (
        START_KEYWORD_REGEX,
        END_KEYWORD_PATTERNS,
        STRONG_DELETE_CONTAIN,
        FORCE_DELETE_PREFIXES,
        ANSWER_VAL_PATTERN
    )

DEBUG_MODE = False

if DEBUG_MODE:
    from docx.enum.text import WD_COLOR_INDEX

def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, Table):
        # Tables don't have direct children like body, need to iterate rows/cells
        # But here we are given a parent, we want to iterate ITS children block items.
        # If parent is Table, it doesn't have block children directly, it has rows.
        # But usually iter_block_items is called on Document or Cell.
        # If called on Table, we should probably yield nothing or raise error.
        # The recursive logic below handles Table by yielding the Table object itself.
        # The main loop handles Table by iterating its cells.
        # So this function is mainly for "Container" objects (Body, Cell).
        return []
    else:
        raise ValueError("Parent object error")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def apply_delete(paragraph, keep_text=None, deleted_text_hint=None, reason=""):
    """
    Apply deletion or highlighting based on DEBUG_MODE
    """
    if DEBUG_MODE:
        reason_str = f" [{reason}]" if reason else " [DEL]"
        
        if keep_text is not None:
            # Partial deletion
            paragraph.text = ""
            paragraph.add_run(keep_text)
            txt = deleted_text_hint if deleted_text_hint else ""
            run = paragraph.add_run(txt + reason_str)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            # Full deletion
            if not paragraph.text.strip():
                paragraph.add_run(f"[DEL IMAGE/EMPTY: {reason}]").font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                for run in paragraph.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                paragraph.add_run(reason_str).font.highlight_color = WD_COLOR_INDEX.YELLOW
    else:
        if keep_text is not None:
            paragraph.text = keep_text
        else:
            paragraph.text = ""

# --- Helper for Validation ---
def check_is_valid_next(found_num, last_num):
    """
    Check if found number is a valid next question number.
    Strictly increasing with SMALL GAP (<= 2).
    """
    if last_num == 0:
        # First question ever: Accept reasonable range (e.g. < 200)
        return 0 < found_num < 200
    
    # Subsequent questions
    diff = found_num - last_num
    return 0 < diff <= 2

def has_image(paragraph):
    """
    Check if a paragraph contains an image (drawing or pict)
    """
    xml = paragraph._element.xml
    return 'w:drawing' in xml or 'w:pict' in xml

def clean_docx_block(input_path, output_path):
    print(f"Processing (Block Mode): {os.path.basename(input_path)}")
    doc = Document(input_path)
    
    if len(doc.paragraphs) > 0:
        title_p = doc.paragraphs[0]
        if "解析" in title_p.text:
            title_p.text = title_p.text.replace("解析", "")

    # State Machine
    is_deleting = False
    last_q_num = 0
    extracted_answers = {}  # {q_num: answer_letter}
    
    blocks = list(iter_block_items(doc))
    
    for block in blocks:
        # Flatten blocks to paragraphs
        paragraphs = []
        if isinstance(block, Paragraph):
            paragraphs.append(block)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        
        for p in paragraphs:
            text = p.text.strip()
            
            # --- EXTRACT ANSWER ---
            # Check for answer pattern even if deleting, because sometimes it's IN the deleted block
            ans_match = ANSWER_VAL_PATTERN.search(text)
            if ans_match and last_q_num > 0:
                # Only store if we haven't found an answer for this Q (or overwrite? usually first match is good)
                extracted_answers[last_q_num] = ans_match.group(1).upper()

            # === STATE: DELETE ===
            if is_deleting:
                # 1. Check STOP Triggers
                stop_match = False
                stop_reason = ""
                
                # Check Question Number
                q_pattern = END_KEYWORD_PATTERNS[0]
                match = q_pattern.match(text)
                if match:
                    try:
                        found_num = int(match.group(1))
                        # Strictly increasing check with SMALL GAP
                        if check_is_valid_next(found_num, last_q_num):
                            stop_match = True
                            stop_reason = f"End: Found Q{found_num}"
                            last_q_num = found_num
                            is_deleting = False
                    except ValueError:
                        pass
                
                # Check Other Headers
                if not stop_match:
                    for pattern in END_KEYWORD_PATTERNS[1:]:
                        if pattern.match(text):
                            stop_match = True
                            stop_reason = "End: Header"
                            is_deleting = False
                            break
                
                if stop_match:
                    # Stopped deleting. Keep this line.
                    if DEBUG_MODE:
                        # Add green marker for the STOP line
                        p.add_run(f" [{stop_reason}]").font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    pass
                else:
                    # No stop trigger -> DELETE EVERYTHING
                    # Don't add text reason for middle blocks (too noisy)
                    apply_delete(p, reason="")
                    continue

            # === STATE: KEEP ===
            if not is_deleting:
                # 1. Check START Trigger
                start_match = START_KEYWORD_REGEX.search(text)
                if start_match:
                    is_deleting = True
                    start_idx = start_match.start()
                    
                    if start_idx > 0:
                        keep = text[:start_idx].strip()
                        deleted = text[start_idx:]
                        # Start line: explicit reason
                        apply_delete(p, keep_text=keep, deleted_text_hint=deleted, reason="Start: Keyword")
                    else:
                        # Start line: explicit reason
                        apply_delete(p, reason="Start: Keyword")
                    
                    # Next iteration will be in DELETE state
                    continue
                
                # 2. Check Force Delete (Single Line)
                is_force = False
                for kw in STRONG_DELETE_CONTAIN:
                    if kw in text:
                        # Also check if this line contains the answer before deleting!
                        # (Already done at top of loop)
                        is_force = True
                        break
                if not is_force:
                    for prefix in FORCE_DELETE_PREFIXES:
                        if text.startswith(prefix):
                            is_force = True
                            break
                
                if is_force:
                    apply_delete(p, reason="Force Delete")
                    continue
                
                # 3. Update Question Number Context (if strictly increasing)
                q_pattern = END_KEYWORD_PATTERNS[0]
                match = q_pattern.match(text)
                if match:
                    try:
                        found_num = int(match.group(1))
                        # ALSO APPLY STRICT CHECK HERE!
                        # Prevents "14237" appearing in Material from corrupting last_q_num
                        if check_is_valid_next(found_num, last_q_num):
                            last_q_num = found_num
                    except ValueError:
                        pass

    # Post Processing - Delete Empty Paragraphs
    count_deleted = 0
    to_delete = []
    for p in doc.paragraphs:
        if not p.text.strip():
            # In DEBUG_MODE, deleted lines have text "[DEL...]", so they won't be empty.
            # Only truly empty lines (whitespace) will be removed.
            
            # PROTECT IMAGES: Check if paragraph has image even if text is empty
            if has_image(p):
                # Do NOT delete paragraph with image
                pass 
            else:
                to_delete.append(p)

    for p in to_delete:
        p_element = p._element
        if p_element.getparent() is not None:
            p_element.getparent().remove(p_element)
            count_deleted += 1

    # --- APPEND EXTRACTED ANSWERS ---
    if extracted_answers:
        doc.add_page_break()
        try:
            doc.add_heading("参考答案", level=1)
        except Exception:
            # Fallback if "Heading 1" style doesn't exist
            p = doc.add_paragraph()
            run = p.add_run("参考答案")
            run.bold = True
            try:
                run.font.size = 20 * 12700 
            except:
                pass
        
        # Sort answers by question number
        sorted_q_nums = sorted(extracted_answers.keys())
        
        # New Layout: Horizontal Table
        # Chunks of 5 seems standard and readable.
        chunk_size = 5
        
        # We need a table with 'chunk_size' columns.
        table = doc.add_table(rows=0, cols=chunk_size)
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
            
        for i in range(0, len(sorted_q_nums), chunk_size):
            chunk = sorted_q_nums[i : i + chunk_size]
            
            # Add Question Row (Top)
            row_q = table.add_row().cells
            # Add Answer Row (Bottom)
            row_a = table.add_row().cells
            
            for j, q_num in enumerate(chunk):
                # Question Number
                row_q[j].text = str(q_num)
                # Answer
                row_a[j].text = extracted_answers[q_num]
                
                # Optional: Center Alignment (if we imported WD_ALIGN_PARAGRAPH)
                # For now just default top-left is safer without new imports.
                
            # If chunk is smaller than chunk_size (last row), fill remaining with empty?
            # docx handles it fine, cells are already created blank.

    doc.save(output_path)
    print(f"  -> Done. Deleted empty paragraphs: {count_deleted}. Extracted {len(extracted_answers)} answers.")

if __name__ == "__main__":
    folder = "."
    files = [f for f in os.listdir(folder) if '解析' in f and f.endswith('.docx')]
    if not files:
        print("No matches found.")
    else:
        print(f"Found {len(files)} files.")
        for f in files:
            in_path = os.path.join(folder, f)
            out_name = f.replace("-解析", "").replace("解析", "")
            if out_name == f: out_name = "题目版_" + f
            out_path = os.path.join(folder, out_name)
            try:
                clean_docx_block(in_path, out_path)
            except Exception as e:
                print(f"Error {f}: {e}")
                import traceback
                traceback.print_exc()
