from fastapi import FastAPI, UploadFile, File, HTTPException, Body
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
import shutil
import os
import uvicorn
from datetime import datetime
from typing import List, Optional, Dict
from pydantic import BaseModel

from extractor import QuestionExtractor
from database import DatabaseManager

from contextlib import asynccontextmanager

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup: Clean temp dir
    temp_dir = os.path.join(MEDIA_DIR, "temp")
    if os.path.exists(temp_dir):
        try:
            shutil.rmtree(temp_dir)
            os.makedirs(temp_dir)
            print("Cleaned media/temp directory.")
        except Exception as e:
            print(f"Failed to clean temp dir: {e}")
    else:
        os.makedirs(temp_dir)
    
    yield
    # Shutdown logic if any

app = FastAPI(lifespan=lifespan)

# ... imports
# Config
import threading
import time
import webbrowser

# ... imports
# Config
import sys



if getattr(sys, 'frozen', False):
    # Running as compiled exe
    # ASSET_DIR: Temporary folder where PyInstaller extracts code/static (Bundle)
    ASSET_DIR = sys._MEIPASS
    # DATA_DIR: Directory where the executable/script resides (User Data)
    DATA_DIR = os.path.dirname(sys.executable)
else:
    # Running as script
    ASSET_DIR = os.path.dirname(os.path.abspath(__file__))
    DATA_DIR = ASSET_DIR

# Mutable User Data (External)
MEDIA_DIR = os.path.join(DATA_DIR, "media")
UPLOAD_DIR = os.path.join(DATA_DIR, "uploads")

if not os.path.exists(UPLOAD_DIR): os.makedirs(UPLOAD_DIR)
if not os.path.exists(MEDIA_DIR): os.makedirs(MEDIA_DIR)

# Init Components
db = DatabaseManager(os.path.join(DATA_DIR, "reservoir.db"))
extractor = QuestionExtractor(MEDIA_DIR)


# ... existing imports ...

class GenerateRequest(BaseModel):
    total_count: int
    types: List[str] = [] # e.g. ["常识", "言语"]

@app.post("/generate")
async def generate_paper(req: GenerateRequest):
    # 1. Get Questions
    if not req.types:
        questions = db.get_standard_exam_questions()
    else:
        questions = db.get_random_questions(req.types, req.total_count)
    
    if not questions:
        raise HTTPException(status_code=404, detail="No questions found matching criteria")

    # 2. Generate UUID & Record
    import uuid
    paper_uuid = str(uuid.uuid4())
    qids = [q['id'] for q in questions]
    db.record_generated_paper(paper_uuid, qids)

    # 3. Generate Doc
    generated_files = create_paper_files(questions, paper_uuid)
    
    # Zip them
    import zipfile
    zip_filename = f"Paper_{datetime.now().strftime('%Y%m%d%H%M%S')}.zip"
    zip_path = os.path.join("media/temp", zip_filename)
    
    with zipfile.ZipFile(zip_path, 'w') as zf:
        for fpath in generated_files:
            if os.path.exists(fpath):
                zf.write(fpath, os.path.basename(fpath))
                
    return FileResponse(zip_path, filename=zip_filename)

def create_paper_files(questions, paper_uuid):
    from generator import PaperBuilder
    generator = PaperBuilder(MEDIA_DIR)
    filename_base = f"Paper_{paper_uuid}.docx"
    output_path_base = os.path.join(MEDIA_DIR, "temp", filename_base)
    os.makedirs(os.path.join(MEDIA_DIR, "temp"), exist_ok=True)
    
    # Process images for docx
    for q in questions:
        # Material Images
        if q.get('material_images'):
            try:
                m_imgs = json.loads(q['material_images']) if isinstance(q['material_images'], str) else q['material_images']
                # Prepend media_dir
                q['material_images_abs'] = [os.path.join(MEDIA_DIR, img) for img in m_imgs]
            except: pass
            
        if q.get('images'):
            try:
                imgs = json.loads(q['images']) if isinstance(q['images'], str) else q['images']
                q['images_abs'] = [os.path.join(MEDIA_DIR, img) for img in imgs]
            except: pass

    return generator.create_paper(questions, output_path_base, paper_uuid=paper_uuid)

@app.get("/api/papers")
def get_paper_history():
    return db.get_all_generated_papers()

@app.get("/api/paper/{uuid}/download")
def download_paper(uuid: str):
    # 1. Get QIDs
    qids = db.get_generated_paper_qids(uuid)
    if not qids:
        raise HTTPException(status_code=404, detail="Paper not found")

    # 2. Fetch Questions (Reusing logic from analyze_file but need a helper or direct query)
    import sqlite3
    conn = db.get_connection()
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    
    query = f'''
        SELECT q.*, s.filename as source_filename, m.content_html as material_content, m.images as material_images
        FROM questions q
        LEFT JOIN sources s ON q.source_id = s.id
        LEFT JOIN materials m ON q.material_id = m.id
        WHERE q.id IN ({','.join(['?']*len(qids))})
    '''
    c.execute(query, qids)
    rows = c.fetchall()
    
    questions_data = []
    for row in rows:
        q = dict(row)
        if q.get('images'): 
            try: q['images'] = json.loads(q['images'])
            except: q['images'] = []
        if q.get('material_images'):
             try: q['material_images'] = json.loads(q['material_images'])
             except: q['material_images'] = []
        questions_data.append(q)
    conn.close()

    # Sort
    questions_map = {q['id']: q for q in questions_data}
    sorted_questions = []
    for i, qid in enumerate(qids):
        if qid in questions_map:
            q = questions_map[qid]
            q['original_num'] = i + 1
            q['num'] = i + 1
            sorted_questions.append(q)

    # 3. Generate
    generated_files = create_paper_files(sorted_questions, uuid)
    
    # 4. Zip
    import zipfile
    zip_filename = f"Paper_{uuid}.zip"
    zip_path = os.path.join(MEDIA_DIR, "temp", zip_filename)
    
    with zipfile.ZipFile(zip_path, 'w') as zf:
        for fpath in generated_files:
            if os.path.exists(fpath):
                zf.write(fpath, os.path.basename(fpath))
                
    return FileResponse(zip_path, filename=zip_filename)

# Mount Static
app.mount("/static", StaticFiles(directory=os.path.join(ASSET_DIR, "static")), name="static")
app.mount("/media", StaticFiles(directory=MEDIA_DIR), name="media")

# Models
class AnalyzeRequest(BaseModel):
    filename: str

class ExtractRequest(BaseModel):
    filename: str
    ranges: Optional[str] = None
    ids: Optional[List[int]] = None

class SaveRequest(BaseModel):
    source_filename: str
    questions: List[dict]
    all_questions_meta: List[dict] # {num, type} for stats calculation
    paper_uuid: Optional[str] = None # Added for review mode
    time_used: Optional[int] = 0 # Added for exam record

class UpdateRequest(BaseModel):
    id: int
    content_html: str
    options_html: str
    answer_html: str

@app.post("/api/question/update")
def update_question(req: UpdateRequest):
    try:
        db.update_question_text(req.id, req.content_html, req.options_html, req.answer_html)
        return {"status": "success"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/question/{qid}")
def delete_question(qid: int):
    try:
        # Get Question Info first to find images
        # We need a db method for single question or search
        # Since db.get_all_questions is heavy, let's just use it or add a fetcher.
        # But for now, let's look at get_all_questions usage.
        # Actually, let's implement a direct fetch in the endpoint using raw cursor for efficiency or add to DB.
        # For simplicity/speed without changing DB schema, we can query.
        
        conn = db.get_connection()
        c = conn.cursor()
        c.execute("SELECT images FROM questions WHERE id=?", (qid,))
        row = c.fetchone()
        images_to_delete = []
        if row and row[0]:
            try:
                images_to_delete = json.loads(row[0])
            except: pass
        conn.close()

        # Delete FileSystem Images
        for img in images_to_delete:
            p = os.path.join(MEDIA_DIR, img)
            if os.path.exists(p):
                try:
                    os.remove(p)
                except Exception as ex:
                    print(f"Failed to delete {p}: {ex}")

        # Delete DB Record
        db.delete_question(qid)
        return {"status": "success"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/exam_stats")
def get_exam_stats():
    try:
        stats = db.get_exam_stats()
        return {"count": len(stats), "stats": stats}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/")
def read_root():
    return FileResponse(os.path.join(ASSET_DIR, "static/index.html"))

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    return {"filename": file.filename}

def parse_ranges(range_str: str) -> List[int]:
    if not range_str: return []
    ids = set()
    parts = range_str.split(',')
    for p in parts:
        p = p.strip()
        if not p: continue
        if '-' in p:
            try:
                start, end = map(int, p.split('-'))
                for i in range(start, end + 1):
                    ids.add(i)
            except: pass
        else:
            try:
                ids.add(int(p))
            except: pass
    return sorted(list(ids))

@app.post("/analyze_file")
def analyze_file(req: AnalyzeRequest):
    file_path = os.path.join(UPLOAD_DIR, req.filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    # --- 1. Detect Paper ID (Review Mode) ---
    paper_uuid = None
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        if len(doc.paragraphs) > 0:
            first_para = doc.paragraphs[0].text.strip()
            if first_para.startswith("Paper ID: "):
                paper_uuid = first_para.replace("Paper ID: ", "").strip()
    except Exception as e:
        print(f"Error checking Paper ID: {e}")
        # Continue to standard extraction if fails (might be PDF or other format)

    if paper_uuid:
        # --- REVIEW MODE (Manual Selection) ---
        # Instead of auto-processing, we fetch the questions and return them to the FE
        # so the user can SELECT which ones they got wrong.
        try:
            qids = db.get_generated_paper_qids(paper_uuid)
            if not qids:
                 return JSONResponse(status_code=404, content={"message": f"Paper ID {paper_uuid} not found locally."})

            # Fetch existing questions
            # We want them to look like "imported" questions but with IDs preserved
            import sqlite3
            conn = db.get_connection()
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            
            # Use a bespoke query to get material content as well, similar to extractor
            # We need standard fields: content_html, type, etc.
            # And we need to ensure images are parsed if they are stored as JSON strings
            
            query = f'''
                SELECT q.*, s.filename as source_filename, m.content_html as material_content
                FROM questions q
                LEFT JOIN sources s ON q.source_id = s.id
                LEFT JOIN materials m ON q.material_id = m.id
                WHERE q.id IN ({','.join(['?']*len(qids))})
            '''
            c.execute(query, qids)
            rows = c.fetchall()
            
            questions_data = []
            for row in rows:
                q = dict(row)
                if q.get('images'): 
                    try:
                        q['images'] = json.loads(q['images'])
                    except:
                        q['images'] = []
                questions_data.append(q)
            conn.close()

            # Re-sort questions to match the order in 'qids' (the generated paper order)
            questions_map = {q['id']: q for q in questions_data}
            sorted_questions = []
            for i, qid in enumerate(qids):
                if qid in questions_map:
                    q = questions_map[qid]
                    # Overwrite numbering to be sequential (1, 2, 3...) for the Review Grid
                    # This matches the "Question 1, Question 2" user sees in the uploaded DOCX
                    q['original_num'] = i + 1
                    q['num'] = i + 1
                    sorted_questions.append(q)

            return {
                "type": "review_import", # New type to signal frontend
                "data": {
                    "paper_uuid": paper_uuid,
                    "questions": sorted_questions
                }
            }
        except Exception as e:
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"Review processing failed: {str(e)}")


    # --- 2. Standard Extraction Mode (Import) ---
    try:
        from extractor import QuestionExtractor
        extractor = QuestionExtractor(MEDIA_DIR)
        questions = extractor.extract_from_file(file_path)
        
        return {
            "type": "import",
            "data": questions
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
        

@app.post("/extract_preview")
def extract_preview(req: ExtractRequest):
    file_path = os.path.join(UPLOAD_DIR, req.filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    target_ids = []
    if req.ids:
        target_ids = req.ids
    elif req.ranges:
        target_ids = parse_ranges(req.ranges)
    
    try:
        # Use temp dir for preview images to prevent zombie files
        questions = extractor.extract_from_file(file_path, target_ids if target_ids else None, sub_dir="temp")
        
        if (req.ids is not None) and len(req.ids) == 0:
             questions = []
             
        return {"count": len(questions), "questions": questions}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))




@app.post("/confirm_save")
def confirm_save(req: SaveRequest):
    new_count = 0
    repeat_count = 0
    count = 0
    
    # 1. Review Mode Branch
    if req.paper_uuid:
        all_paper_qids = db.get_generated_paper_qids(req.paper_uuid)
        if not all_paper_qids:
             return {"status": "error", "message": "Paper not found"}
        
        wrong_qids = [q['id'] for q in req.questions if q.get('id')]
        db.process_review_results(wrong_qids, all_paper_qids)
        count = len(wrong_qids)
        
    # 2. Import Mode Branch
    else:
        try:
            # 1. Add Source (Normal Import)
            sid = db.add_source(req.source_filename)
            
            # 2. Add Questions & Materials
            material_map = {} # content_hash -> mid
            
            # Helper to move file if exists
            def move_from_temp(filename):
                src = os.path.join(MEDIA_DIR, "temp", filename)
                dst = os.path.join(MEDIA_DIR, filename)
                if os.path.join(MEDIA_DIR, "temp") in src and os.path.exists(src):
                    try:
                        shutil.move(src, dst)
                    except Exception as e:
                        print(f"Error moving {filename}: {e}")

            # Helper to fix HTML
            def fix_html_paths(html):
                if not html: return html
                return html.replace("/media/temp/", "/media/")

            for q in req.questions:
                # Move physical files
                if q.get('images'):
                    for img in q['images']:
                        move_from_temp(img)
                
                # Update HTML
                q['content_html'] = fix_html_paths(q['content_html'])
                q['options_html'] = fix_html_paths(q['options_html'])
                q['answer_html'] = fix_html_paths(q['answer_html'])

                # Material handling
                mid = None
                mat_content = q.get('material_content')
                if mat_content:
                    import re
                    mat_temp_imgs = re.findall(r'/media/temp/([\w\-\.]+\.\w+)', mat_content)
                    for img in mat_temp_imgs:
                        move_from_temp(img)
                    
                    mat_content = fix_html_paths(mat_content)
                    
                    mat_hash = hash(mat_content)
                    if mat_hash in material_map:
                        mid = material_map[mat_hash]
                    else:
                        mid = db.add_material(sid, mat_content, type=q['type'])
                        material_map[mat_hash] = mid
                
                # Check for missing keys or defaults
                q_type = q.get('type', 'Unknown')
                q_images = q.get('images', [])

                qid, is_new = db.add_question(
                    source_id=sid,
                    original_num=q['original_num'],
                    content=q['content_html'],
                    options=q['options_html'],
                    answer=q['answer_html'], 
                    images=q_images,
                    type=q_type,
                    material_id=mid
                )
                
                if is_new:
                    new_count += 1
                else:
                    repeat_count += 1
                    
                count += 1
        except Exception as e:
            import traceback
            traceback.print_exc()
            return {"status": "error", "message": f"Save failed details: {str(e)}"}
            
    # --- Stats Calculation (Shared) ---
    try:
        # Weights
        WEIGHTS = {
            "常识": 0.5,
            "言语": 0.8,
            "数量": 0.8,
            "判断": 0.7,
            "资料": 1.0,
            # Subtypes just in case
            "图形": 0.7, "定义": 0.7, "类比": 0.7, "逻辑": 0.7
        }

        mistake_nums = set(str(q.get('original_num', '')) for q in req.questions)
        
        module_stats = {} # {"Verbal": {correct: 0, total: 0}}
        total_score = 0.0
        total_questions = 0
        total_correct = 0

        for q in req.all_questions_meta:
            q_type = q.get('type', '未知')
            
            # Find Standard Module Name (grouping subtypes)
            module_name = "未知"
            if "常识" in q_type: module_name = "常识"
            elif "言语" in q_type: module_name = "言语"
            elif "数量" in q_type: module_name = "数量"
            elif "资料" in q_type: module_name = "资料"
            elif any(x in q_type for x in ["判断", "图形", "定义", "类比", "逻辑"]): module_name = "判断"

            # Init Module Entry
            if module_name not in module_stats:
                module_stats[module_name] = {"correct": 0, "total": 0}
            
            # Safe access to num, check 'num' then 'original_num'
            # Frontend might pass extractor output directly which uses 'original_num'
            val = q.get('num')
            if val is None:
                val = q.get('original_num')
            
            q_num = str(val if val is not None else -1)
            is_mistake = q_num in mistake_nums
            
            # Debug (Temporary, remove if too spammy)
            # print(f"Checking Q {q_num}: Mistake? {is_mistake} (in {mistake_nums})")
            
            # Update Total
            module_stats[module_name]["total"] += 1
            total_questions += 1
            
            # Update Correct & Score
            if not is_mistake:
                module_stats[module_name]["correct"] += 1
                total_correct += 1
                
                # Add Score
                w = WEIGHTS.get(module_name, 0.5) # Default 0.5 for unknown
                total_score += w

        total_accuracy = (total_correct / total_questions * 100) if total_questions > 0 else 0.0
        
        # Determine filename for record
        record_name = req.source_filename
        if req.paper_uuid:
             # Use a distinct name or the original filename with marker
             # Reusing source_filename usually works if frontend sends it, 
             # but let's make it clear it's a review.
             record_name = f"Review_{req.paper_uuid[:8]}"

        db.add_exam_record(record_name, total_score, total_accuracy, module_stats, time_used=req.time_used)
        
    except Exception as e:
        print(f"Stats Calculation Failed: {e}")

        # Non-blocking, still return success for saving

    return {"status": "success", "saved_count": count, "new_count": new_count, "repeat_count": repeat_count}



@app.get("/pool_status")
def pool_status():
    return db.get_pool_status()


@app.get("/api/questions")
def get_all_questions():
    questions = db.get_all_questions()
    return {"count": len(questions), "questions": questions}

@app.get("/browse")
def browse_page():
    return FileResponse(os.path.join(ASSET_DIR, "static/browse.html"))

# To run: uvicorn main:app --reload
def find_available_port(start_port, max_port=65535):
    import socket
    for port in range(start_port, max_port):
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            try:
                s.bind(("127.0.0.1", port))
                return port
            except OSError:
                continue
    return None

# System Tray & Server Startup
def run_server(port):
    # Redirect streams for no-console mode
    if sys.stdout is None: sys.stdout = open(os.devnull, "w")
    if sys.stderr is None: sys.stderr = open(os.devnull, "w")
    if sys.stdin is None: sys.stdin = open(os.devnull, "r")
    
    try:
        uvicorn.run(app, host="127.0.0.1", port=port, log_level="error")
    except Exception as e:
        # Minimal fail-safe error log
        if getattr(sys, 'frozen', False):
             try:
                 with open("startup_error.log", "w") as f:
                     f.write(f"Startup Failed: {str(e)}\n")
             except: pass

def setup_tray(url):
    try:
        import pystray
        from PIL import Image
    except ImportError:
        print("pystray or PIL not installed. Running without tray.")
        return None

    def on_open(icon, item):
        webbrowser.open(url)

    def on_quit(icon, item):
        icon.stop()
        os._exit(0)

    # Load Icon
    icon_path = os.path.join(ASSET_DIR, "approved.png")
    if not os.path.exists(icon_path):
        # Fallback to creating a simple image if missing
        image = Image.new('RGB', (64, 64), color = (73, 109, 137))
    else:
        image = Image.open(icon_path)

    menu = pystray.Menu(
        pystray.MenuItem("Open MistakeReservoir", on_open, default=True),
        pystray.MenuItem("Quit", on_quit)
    )

    icon = pystray.Icon("MistakeReservoir", image, "MistakeReservoir", menu)
    return icon

if __name__ == "__main__":
    import argparse
    import threading
    import webbrowser
    import time
    import sys
    import os
    
    # Simple CLI Args
    parser = argparse.ArgumentParser(description="MistakeReservoir Server")
    parser.add_argument("--t", action="store_true", help="Enter terminal debug mode (no server)")
    parser.add_argument("--port", type=int, default=8000, help="Starting port number")
    args = parser.parse_args()
    
    # Debug Mode
    if args.t:
        print("Entering Terminal Debug Mode...")
        print("Variables available: app, db, extractor, etc.")
        import code
        context = globals().copy()
        context.update(locals())
        code.interact(local=context, banner="MistakeReservoir Debug Shell")
        sys.exit(0)

    # Server Mode
    port = find_available_port(args.port)
    if not port:
        if sys.stdout: print("Error: No available ports found.")
        sys.exit(1)
        
    url = f"http://127.0.0.1:{port}"
    print(f"Starting server on port {port}...")
    
    # Strat Server in Thread
    server_thread = threading.Thread(target=run_server, args=(port,), daemon=True)
    server_thread.start()
    
    # Open Browser
    def open_browser():
        time.sleep(1.5)
        webbrowser.open(url)
    threading.Thread(target=open_browser, daemon=True).start()
    
    # Run Tray (Main Thread)
    tray = setup_tray(url)
    if tray:
        tray.run()
    else:
        # Fallback if no tray lib
        try:
            while True: time.sleep(1)
        except KeyboardInterrupt:
            os._exit(0)
