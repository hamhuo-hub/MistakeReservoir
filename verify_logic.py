import os
import uuid
import json
import sqlite3
from docx import Document
from main import app, analyze_file, AnalyzeRequest, UPLOAD_DIR
from main import db # Import db from main to ensure initialized state? or re-init?
# DatabaseManager is in database.py
from database import DatabaseManager
import sys

# Mocking FastAPI dependency if needed, but we invoke function directly.
# However, analyze_file uses 'db' global in main.py. 
# We need to make sure 'db' in main.py is accessible or we rely on it being there.

# The script imports 'analyze_file' from 'main'. 'main' initializes 'db'.
# So 'db' inside 'analyze_file' should work. 

# But we need 'db' here to record paper.
db_local = DatabaseManager()

# Setup
db = DatabaseManager()
test_uuid = str(uuid.uuid4())
print(f"Test UUID: {test_uuid}")

# 1. Mock Database Record
# Get some real question IDs to link
conn = db.get_connection()
c = conn.cursor()
c.execute("SELECT id FROM questions LIMIT 5")
qids = [r[0] for r in c.fetchall()]
conn.close()

if not qids:
    print("Skipping test: No questions in DB.")
    exit()

print(f"Linking UUID to QIDs: {qids}")
db_local.record_generated_paper(test_uuid, qids)

# 2. Create Dummy Review File (DOCX)
review_filename = f"test_review_{test_uuid}.docx"
review_path = os.path.join(UPLOAD_DIR, review_filename)
doc = Document()
doc.add_paragraph(f"Paper ID: {test_uuid}")
doc.add_paragraph("Some question content here that doesn't matter for ID detection but matters for matching.")
doc.save(review_path)

# 3. Test Analyze - Review Mode
print(f"\n--- Testing Review Mode with {review_filename} ---")
try:
    res = analyze_file(AnalyzeRequest(filename=review_filename))
    if isinstance(res, dict) and res.get('type') == 'review':
        print("✅ Correctly detected REVIEW mode.")
        print(f"Data: {res['data']['paper_id']}")
    else:
        print(f"❌ Failed. Expected 'review', got: {res}")
except Exception as e:
    print(f"❌ Exception: {e}")

# 4. Create Dummy Import File (No ID)
import_filename = f"test_import_{test_uuid}.docx"
import_path = os.path.join(UPLOAD_DIR, import_filename)
doc = Document()
doc.add_paragraph("Just some questions.")
doc.add_paragraph("1. What is 1+1?")
doc.add_paragraph("A. 1 B. 2")
doc.save(import_path)

# 5. Test Analyze - Import Mode
print(f"\n--- Testing Import Mode with {import_filename} ---")
try:
    # Note: simple text might yield 0 questions but should still return type='import'
    res = analyze_file(AnalyzeRequest(filename=import_filename))
    if isinstance(res, dict) and res.get('type') == 'import':
        print("✅ Correctly detected IMPORT mode.")
    else:
        print(f"❌ Failed. Expected 'import', got: {res}")
except Exception as e:
    print(f"❌ Exception: {e}")

# Cleanup
try:
    os.remove(review_path)
    os.remove(import_path)
except: pass
