import os
import shutil
from docx import Document
from extractor import QuestionExtractor

def create_dummy_docx(filename):
    doc = Document()
    doc.add_heading('第一部分 常识判断', level=1)
    
    # Q1
    doc.add_paragraph('1. This is the stem of question 1.')
    doc.add_paragraph('A. Option A')
    doc.add_paragraph('B. Option B')
    doc.add_paragraph('C. Option C')
    doc.add_paragraph('D. Option D')
    doc.add_paragraph('【答案】A')
    doc.add_paragraph('【解析】This is the analysis.')
    
    # Material Q
    doc.add_paragraph('根据以下材料回答2-3题')
    doc.add_paragraph('Material Content Here.')
    
    # Q2
    doc.add_paragraph('2. Question 2 Stem.')
    doc.add_paragraph('A. Op A')
    doc.add_paragraph('B. Op B')
    doc.add_paragraph('C. Op C')
    doc.add_paragraph('D. Op D')
    doc.add_paragraph('【答案】B')
    doc.add_paragraph('【解析】Analysis 2.')
    
    doc.save(filename)
    return filename

def test_extraction():
    test_file = "test_data.docx"
    create_dummy_docx(test_file)
    
    # Setup media dir
    media_dir = "test_media"
    if os.path.exists(media_dir):
        shutil.rmtree(media_dir)
    os.makedirs(media_dir)
    
    try:
        extractor = QuestionExtractor(media_dir)
        questions = extractor.extract_from_file(test_file)
        
        print(f"Extracted {len(questions)} questions.")
        
        # Verify Q1
        q1 = questions[0]
        assert q1['original_num'] == 1
        assert "This is the stem of question 1" in q1['content_html']
        assert "Option A" in q1['options_html']
        assert "A" in q1['answer_html'] 
        assert "This is the analysis" in q1['answer_html']
        assert q1['type'] == "常识"
        print("Q1 Verified.")
        
        # Verify Q2 (Material)
        q2 = questions[1]
        assert q2['original_num'] == 2
        assert "Material Content Here" in q2['material_content']
        print("Q2 Verified.")
        
        print("All tests passed!")
        
    except Exception as e:
        print(f"Test Failed: {e}")
        import traceback
        traceback.print_exc()
    finally:
        # Cleanup
        if os.path.exists(test_file):
            os.remove(test_file)
        if os.path.exists(media_dir):
            shutil.rmtree(media_dir)

if __name__ == "__main__":
    test_extraction()
